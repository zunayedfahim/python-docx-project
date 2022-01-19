from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

end_electric_reading_date = str(input('Enter Billing Date:'))
start_electric_reading_date = str(input('Enter Start Electric Reading Date:'))

flats = ['A-1', 'A-2', 'B-2', 'A-3', 'B-3', 'A-4',
         'B-4', 'A-5', 'B-5', 'A-6', 'B-6', 'A-7', 'B-7']

total_taka = 0



# Count bill using the given consumed unit.
def count_bill(unit):
    if unit == 0:
        return 0
    bill = 0

    if 0 < unit <= 75:
        bill = unit*4.19

    elif 76 <= unit <= 200:
        bill = 314.25 + ((unit-75)*5.72)

    elif 201 <= unit <= 300:
        bill = 1029.25 + ((unit-200)*6.00)

    elif 301 <= unit <= 400:
        bill = 1629.25 + ((unit-300)*6.34)

    elif 401 <= unit <= 600:
        bill = 2263.25 + ((unit-400)*9.94)

    elif unit > 601:
        bill = 4251.25 + ((unit-600)*11.46)

    else:
        bill = 'wrong input'

    bill_w_dc = bill+90
    bill_w_vat = bill_w_dc + ((bill_w_dc)*0.05)
    return int(bill_w_vat)

# Insert a comma if the value is above thousand.
def insert_comma(word):
    word = str(word)
    if len(word) == 4:
        return word[0] + ',' + word[1:]
    elif len(word) == 5:
        return word[:2] + ',' + word[2:]
    elif len(word) == 6:
        return word[0] + ',' + word[1:3] + ',' + word[3:]
    else:
        return word


document = Document()

sections = document.sections
for section in sections:
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.2)
    section.left_margin = Inches(0.2)
    section.right_margin = Inches(0.2)

# for i in range(1):
for flat_no in flats:
    valid_input = False
    while not valid_input:
        try:
            print('-----Enter for ' + flat_no + ' -----')
            # flat_no = input('Flat Number:')
            billing_date = 'BILLING DATE: ' + end_electric_reading_date

            if flat_no == 'A-1':
                meter_no = '10798704'
                name = 'MR.OMAR FARUK'
                so = 'S/O: LATE NAZRUL ISLAM'
                advance = 'ADVANCE: TK.30,000'
                rent = 10500

            elif flat_no == 'A-2':
                meter_no = '10798696'
                name = 'MD.ENAMUL HOQUE CHOWDHURY'
                so = 'S/O: MD. SARU MIA CHOWDHURY'
                advance = 'ADVANCE: TK.21,000'
                rent = 10500

            elif flat_no == 'B-2':
                meter_no = '10798705'
                name = 'MR. S.K.M ZAHANGIR ALAM'
                so = ' S/O: MR. S.K. IMAN UDDIN'
                advance = 'ADVANCE: TK.20,000'
                rent = 10000

            elif flat_no == 'A-3':
                meter_no = '10798701'
                name = 'MR. MD.IMRAN KHAN'
                so = 'S/O: MR. ABDUL JALIL KHAN'
                advance = 'ADVANCE: TK.21,000'
                rent = 10500

            elif flat_no == 'B-3':
                meter_no = '10798693'
                name = 'MR. ABDUL KARIM ANSARI'
                so = 'S/O: MR. HAFEZ AHMED '
                advance = 'ADVANCE: TK.19,000'
                rent = 9500

            elif flat_no == 'A-4':
                meter_no = '10798702'
                name = 'MR. TAREK HASAN SABBIR'
                so = 'S/O: LATE ABU TAHER'
                advance = 'ADVANCE: TK.20,000'
                rent = 10000

            elif flat_no == 'B-4':
                meter_no = '10808826'
                name = 'MOHAMMAD JABER'
                so = 'S/O: LATE MOULANA FORKAN'
                advance = 'ADVANCE: TK.19,000'
                rent = 9500

            elif flat_no == 'A-5':
                meter_no = '10808828'
                name = 'MR. S.M HAMID ANWAR'
                so = 'S/O: LATE SULTAN AHMED ANWAR'
                advance = 'ADVANCE: TK.28,500'
                rent = 9500

            elif flat_no == 'B-5':
                meter_no = '10808829'
                name = 'MR. MOHAMMAD LOKMAN GANI'
                so = 'S/O: MD. SYEDUL  HOQUE'
                advance = 'ADVANCE: TK.17,000'
                rent = 8500

            elif flat_no == 'A-6':
                meter_no = '10798703'
                name = 'MR. DEWAN SHAHAN GAZI'
                so = 'S/O. DEWAN JAMSHED GAZI'
                advance = 'ADVANCE: TK.18,000'
                rent = 9000

            elif flat_no == 'B-6':
                meter_no = '10808814'
                name = 'MR. ASHOK CHOWDHURY'
                so = 'S/O: LATE NUPUR KANTI CHOWDHURY'
                advance = 'ADVANCE: TK.17,000'
                rent = 8500

            elif flat_no == 'A-7':
                meter_no = '10808831'
                name = 'MR. RIYAD MAHMUD CHOWDHURY'
                so = 'S/O: MOHAMMAD HEFAZATUL ISLAM CHOWDHURY'
                advance = 'ADVANCE: TK.13,000'
                rent = 6500

            elif flat_no == 'B-7':
                meter_no = '10741625'
                name = 'MR. SHARIFUL ISLAM'
                so = 'S/O: MR. RUHUL AMIN'
                advance = 'ADVANCE: TK.15,000'
                rent = 7500

            prev_out = int(input("Previous Outstanding:"))
            PREV_unit = int(input("Previous electric reading:"))
            CURR_unit = int(input("Current electric reading:"))
            unit = CURR_unit - PREV_unit

            electric_bill = count_bill(unit)
            if electric_bill > 0:
                bill_rate = round(electric_bill/unit, 2)
            else:
                bill_rate = 0
            total_rent = electric_bill + prev_out + rent
            total_taka += total_rent

            if prev_out == 0:
                prev_out = 'NIL'

            d1 = document.add_paragraph()
            d1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            d1.paragraph_format.space_after = Pt(0)
            details_1 = d1.add_run(
                'HOUSE RENT & ELECTRIC BILL\nADDRESS: D/17, ROAD NO. 1, KOLPOLOK R/A, CHATTAGRAM')
            d2 = document.add_paragraph()
            d2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            d2.paragraph_format.space_after = Pt(0)
            details_2 = d2.add_run(
                'NAME: ' + name + '\n' + so + '\n' + advance + '\n' + billing_date)
            details_1.bold = True
            details_2.bold = True
            details_1.underline = True
            details_1.font.size = Pt(16)
            details_2.font.size = Pt(12)

            prev_out = insert_comma(prev_out)
            rent = insert_comma(rent)
            total_rent = insert_comma(total_rent)

            records = (
                (str(flat_no), str(meter_no), str(prev_out), str(rent), str(PREV_unit), str(CURR_unit),
                 str(unit) + '\n@' + str(bill_rate), str(electric_bill), str(total_rent)),
            )

            table = document.add_table(rows=1, cols=10, style='Table Grid')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'FLAT NO.'
            hdr_cells[1].text = 'METER NO.'
            hdr_cells[2].text = 'PREVIOUS OUTSTANDING'
            hdr_cells[3].text = 'RENT & PARKING CHARGE'
            hdr_cells[4].text = 'INITIAL READING ' + start_electric_reading_date
            hdr_cells[5].text = 'CURRENT READING ' + end_electric_reading_date
            hdr_cells[6].text = 'CONSUMED UNIT TILL ' + end_electric_reading_date
            hdr_cells[7].text = 'TAKA'
            hdr_cells[8].text = 'RECIEVER\'S SIGNATURE'
            hdr_cells[9].text = 'DUE AMOUNT IN TAKA'
            for flat_number, meter_number, prev_out, rent, start, end, consumed, taka, due in records:
                row_cells = table.add_row().cells
                row_cells[0].text = flat_number
                row_cells[1].text = meter_number
                row_cells[2].text = prev_out
                row_cells[3].text = rent
                row_cells[4].text = start
                row_cells[5].text = end
                row_cells[6].text = consumed
                row_cells[7].text = taka
                row_cells[8].text = ' '
                row_cells[9].text = due

            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(10)

            p = document.add_paragraph()
            valid_input = True
        except Exception as e:
            print("----------------------------Invalid Input. Please try again.")

total_taka = insert_comma(total_taka)
q_ = document.add_paragraph()
q_.alignment = WD_ALIGN_PARAGRAPH.CENTER
q = q_.add_run('TOTAL TAKA = TK. ' + str(total_taka))
q.bold = True
q.font.size = Pt(20)

wasa_bill = input('Close the file if its open. Any WASA motor bill? (y/n):')
if wasa_bill == 'y' or wasa_bill == "Y":
    Prev_unit = int(input("Previous motor reading:"))
    Curr_unit = int(input("Current motor reading:"))
    unit_ = Curr_unit - Prev_unit
    wasa_bill_ = count_bill(unit_)
    wasa_bill_ = insert_comma(wasa_bill_)

    r_ = document.add_paragraph()
    r_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = r_.add_run('Parking and WASA motor bill ('+str(Curr_unit)+"-"+str(Prev_unit)+") = "+str(unit_)+" = TK."+ str(wasa_bill_))
    r.bold = True
    r.font.size = Pt(30)
filename = input("Filename: ")
document.save('D:\\kolpolok\\' + filename +'.docx')
